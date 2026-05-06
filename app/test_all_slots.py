"""
Test script: Generate 5 mock .docx files with deliberate formatting errors,
upload them to the API, and verify hard rules are caught for ALL slots.
"""
import os, io, sys, json, time
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

API = "http://127.0.0.1:8000"
SID = "test-session-001"
PASS = 0
FAIL = 0

def make_slot1():
    """Front Matter with WRONG margins and missing keywords."""
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0)   # Should be 2.5
    sec.left_margin = Inches(1.0)  # Should be 1.5
    sec.right_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    p = doc.add_paragraph("ชื่อวิทยานิพนธ์ทดสอบ")
    run = p.runs[0]
    run.font.name = "Arial"  # Wrong font
    run.font.size = Pt(14)   # Wrong size
    doc.add_paragraph("โดย นายทดสอบ ทดสอบ")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def make_slot2():
    """Abstract with wrong margins and no indentation."""
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0)   # Should be 1.5
    sec.left_margin = Inches(1.0)  # Should be 1.5
    sec.right_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    p = doc.add_paragraph("บทคัดย่อ")
    run = p.runs[0]
    run.font.size = Pt(16)  # Should be 18
    run.font.name = "TH Sarabun New"
    body = doc.add_paragraph(
        "การวิจัยครั้งนี้มีวัตถุประสงค์เพื่อศึกษาปัจจัยที่มีผลต่อ"
        "ความพึงพอใจของลูกค้า โดยใช้แบบสอบถามเป็นเครื่องมือในการเก็บข้อมูล"
    )
    body.paragraph_format.first_line_indent = Inches(0.0)  # Should be 0.5
    body.runs[0].font.name = "TH Sarabun New"
    body.runs[0].font.size = Pt(14)  # Wrong size
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def make_slot3():
    """TOC without Thai page numbering."""
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0)  # Should be 1.5
    sec.left_margin = Inches(1.5)
    sec.right_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    p = doc.add_paragraph("สารบัญ")
    run = p.runs[0]
    run.font.size = Pt(16)  # Should be 18
    run.font.name = "TH Sarabun New"
    for i in range(1, 4):
        entry = doc.add_paragraph(f"บทที่ {i} .............. {i}")
        entry.runs[0].font.name = "TH Sarabun New"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def make_slot4():
    """Main content with wrong chapter title formatting."""
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0)  # Should be 1.5
    sec.left_margin = Inches(1.0) # Should be 1.5
    sec.right_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    p = doc.add_paragraph("บทที่ 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Should be CENTER
    run = p.runs[0]
    run.font.size = Pt(16)  # Should be 18
    run.font.name = "TH Sarabun New"
    body = doc.add_paragraph("บทนำ เนื้อหาตัวอย่าง สำหรับทดสอบระบบการตรวจสอบรูปแบบเอกสาร")
    body.runs[0].font.name = "TH Sarabun New"
    body.runs[0].font.size = Pt(16)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def make_slot5():
    """References with no hanging indent."""
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0)  # Should be 1.5
    sec.left_margin = Inches(1.5)
    sec.right_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    p = doc.add_paragraph("เอกสารอ้างอิง")
    run = p.runs[0]
    run.font.size = Pt(16)  # Should be 18
    run.font.name = "TH Sarabun New"
    refs = [
        "สมชาย ใจดี. (2563). การศึกษาเรื่องทดสอบ. วารสารวิชาการ, 10(2), 100-115.",
        "Smith, J. (2020). A study of testing systems. Journal of Testing, 5(1), 50-65.",
        "Johnson, A., & Brown, B. (2021). Advanced test methods in software engineering. New York: Academic Press.",
    ]
    for ref_text in refs:
        rp = doc.add_paragraph(ref_text)
        rp.paragraph_format.first_line_indent = Inches(0)  # No hanging indent
        rp.runs[0].font.name = "TH Sarabun New"
        rp.runs[0].font.size = Pt(16)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def test_slot(slot_num, file_buf, url_path):
    global PASS, FAIL
    print(f"\n{'='*60}")
    print(f"  TESTING SLOT {slot_num}: {url_path}")
    print(f"{'='*60}")
    try:
        resp = requests.post(
            f"{API}{url_path}?session_id={SID}",
            files={"file": (f"test_slot{slot_num}.docx", file_buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            timeout=30,
        )
        if resp.status_code != 200:
            print(f"  ❌ HTTP {resp.status_code}: {resp.text[:200]}")
            FAIL += 1
            return
        data = resp.json()
        hard_errors = data.get("hard_rules_errors", [])
        err_count = data.get("error_count", 0)
        warn_count = data.get("warning_count", 0)
        status = data.get("status", "unknown")
        print(f"  Status: {status}")
        print(f"  Hard Errors: {err_count} errors, {warn_count} warnings")
        print(f"  Total hard_rules_errors items: {len(hard_errors)}")
        if hard_errors:
            for i, e in enumerate(hard_errors):
                sev = e.get("severity", "?")
                icon = "🔴" if sev == "error" else "⚠️"
                print(f"    {icon} [{e.get('rule_type','')}] {e.get('issue','')} — {e.get('detail','')[:80]}")
        if len(hard_errors) > 0:
            print(f"  ✅ PASS — Slot {slot_num} returned {len(hard_errors)} hard rule findings")
            PASS += 1
        else:
            print(f"  ❌ FAIL — Slot {slot_num} returned 0 hard rule errors (expected >0)")
            FAIL += 1
    except Exception as ex:
        print(f"  ❌ EXCEPTION: {ex}")
        FAIL += 1

if __name__ == "__main__":
    print("\n" + "🚀 " * 20)
    print("  KKU THESIS CHECKER — ALL-SLOT HARD RULES TEST")
    print("🚀 " * 20)
    
    test_slot(1, make_slot1(), "/api/slot/1")
    test_slot(2, make_slot2(), "/api/slot/2")
    test_slot(3, make_slot3(), "/api/slot/3")
    test_slot(4, make_slot4(), "/api/slot/4/1")
    test_slot(5, make_slot5(), "/api/slot/5")
    
    print(f"\n{'='*60}")
    print(f"  FINAL RESULTS: {PASS} PASSED, {FAIL} FAILED out of 5")
    print(f"{'='*60}")
    if FAIL == 0:
        print("  🎉 ALL SLOTS PASS — Hard rules working for every slot!")
    else:
        print(f"  ⚠️ {FAIL} slot(s) failed. Review output above.")
    sys.exit(FAIL)
