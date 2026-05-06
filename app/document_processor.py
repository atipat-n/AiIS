"""
Hard Rules Checker — document_processor.py
Validates margins, fonts, title sizes, indentation, keywords, and logo
in .docx files based on KKU thesis formatting guidelines.

v4: Full hard-rule coverage for ALL 5 slots.
"""

import io
import re
from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models import HardRuleError

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False


# ──────────────────────────────────────────────
# Margin definitions per slot (inches)
# ──────────────────────────────────────────────

SLOT_MARGINS = {
    1: {"Top": 2.5, "Left": 1.5, "Bottom": 1.0, "Right": 1.0},
    2: {"Top": 1.5, "Left": 1.5, "Bottom": 1.0, "Right": 1.0},
    3: {"Top": 1.5, "Left": 1.5, "Bottom": 1.0, "Right": 1.0},
    4: {"Top": 1.5, "Left": 1.5, "Bottom": 1.0, "Right": 1.0},
    5: {"Top": 1.5, "Left": 1.5, "Bottom": 1.0, "Right": 1.0},
}

MARGIN_TOLERANCE = 0.15  # inches


def check_margins(doc: Document, slot_number: int) -> list[HardRuleError]:
    expected = SLOT_MARGINS.get(slot_number)
    if expected is None:
        return []
    errors = []
    for sec_idx, section in enumerate(doc.sections):
        actual = {
            "Top": section.top_margin.inches if section.top_margin else 0,
            "Bottom": section.bottom_margin.inches if section.bottom_margin else 0,
            "Left": section.left_margin.inches if section.left_margin else 0,
            "Right": section.right_margin.inches if section.right_margin else 0,
        }
        for side, expected_val in expected.items():
            actual_val = actual[side]
            if abs(actual_val - expected_val) > MARGIN_TOLERANCE:
                errors.append(HardRuleError(
                    location=f"Section {sec_idx+1}, Page Setup",
                    issue=f"ระยะขอบ{_thai_side(side)}ผิด (Wrong {side} Margin)",
                    detail=(
                        f"พบ {actual_val:.2f} นิ้ว (Found {actual_val:.2f} in). "
                        f"ค่าที่ถูกต้อง: {expected_val:.1f} นิ้ว (Expected {expected_val:.1f} in)"
                    ),
                    severity="error",
                    rule_type="margin",
                ))
    return errors


def _thai_side(side: str) -> str:
    return {"Top": "บน", "Bottom": "ล่าง", "Left": "ซ้าย", "Right": "ขวา"}.get(side, side)


# ──────────────────────────────────────────────
# Font checks (Global)
# ──────────────────────────────────────────────

ALLOWED_FONTS = {"TH Sarabun New", "TH SarabunPSK", "TH Sarabun PSK"}


def check_fonts(doc: Document) -> list[HardRuleError]:
    errors = []
    flagged_paras = set()
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        for run in para.runs:
            font_name = run.font.name
            if font_name is not None and font_name not in ALLOWED_FONTS:
                if i not in flagged_paras:
                    flagged_paras.add(i)
                    preview = para.text.strip()[:60]
                    errors.append(HardRuleError(
                        location=f"Paragraph {i+1}",
                        issue="ฟอนต์ผิด (Wrong Font)",
                        detail=(
                            f"ใช้ฟอนต์ '{font_name}' ในย่อหน้า: \"{preview}...\" "
                            f"(ควรใช้ TH Sarabun New / TH SarabunPSK)"
                        ),
                        severity="error",
                        rule_type="font",
                    ))
    return errors


# ──────────────────────────────────────────────
# Normal text size check (16pt)
# ──────────────────────────────────────────────

def check_normal_text_size(doc: Document) -> list[HardRuleError]:
    errors = []
    flagged = set()
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or len(text) < 5:
            continue
        for run in para.runs:
            sz = run.font.size
            if sz is not None:
                pt_val = sz.pt
                if pt_val not in (16, 18) and i not in flagged:
                    flagged.add(i)
                    preview = text[:50]
                    errors.append(HardRuleError(
                        location=f"Paragraph {i+1}",
                        issue=f"ขนาดตัวอักษรผิด (Wrong Font Size: {pt_val:.0f}pt)",
                        detail=(
                            f"ย่อหน้า: \"{preview}...\" — ขนาด {pt_val:.0f}pt "
                            f"(เนื้อหาปกติควรเป็น 16pt, หัวข้อ 18pt)"
                        ),
                        severity="warning",
                        rule_type="font_size",
                    ))
    return errors


# ──────────────────────────────────────────────
# Title checks (18pt, Bold, Centered)
# ──────────────────────────────────────────────

def check_titles(doc: Document, expected_titles: list[str], slot_label: str) -> list[HardRuleError]:
    """Check that paragraphs matching expected_titles are 18pt, bold, centered."""
    errors = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        is_title = any(t in text for t in expected_titles)
        if not is_title:
            continue
        # Check size
        sizes = [run.font.size.pt for run in para.runs if run.font.size]
        has_18 = any(abs(s - 18) < 0.5 for s in sizes)
        if sizes and not has_18:
            sz_str = ", ".join(f"{s:.0f}pt" for s in sizes)
            errors.append(HardRuleError(
                location=f"Paragraph {i+1} ({slot_label})",
                issue="ขนาดหัวข้อผิด (Title Size ≠ 18pt)",
                detail=f"หัวข้อ \"{text[:50]}\" — พบ {sz_str} (ควรเป็น 18pt Bold)",
                severity="error",
                rule_type="title_size",
            ))
        # Check bold
        has_bold = any(run.font.bold for run in para.runs if run.text.strip())
        if not has_bold:
            errors.append(HardRuleError(
                location=f"Paragraph {i+1} ({slot_label})",
                issue="หัวข้อไม่เป็นตัวหนา (Title Not Bold)",
                detail=f"หัวข้อ \"{text[:50]}\" ควรเป็นตัวหนา (Bold)",
                severity="warning",
                rule_type="title_bold",
            ))
        # Check centered
        alignment = para.alignment
        if alignment is not None and alignment != WD_ALIGN_PARAGRAPH.CENTER:
            errors.append(HardRuleError(
                location=f"Paragraph {i+1} ({slot_label})",
                issue="หัวข้อไม่จัดกึ่งกลาง (Title Not Centered)",
                detail=f"หัวข้อ \"{text[:50]}\" ควรจัดกึ่งกลาง (Center)",
                severity="warning",
                rule_type="title_align",
            ))
    return errors


# ──────────────────────────────────────────────
# Slot 1: Front Matter specific checks
# ──────────────────────────────────────────────

SLOT1_REQUIRED_KEYWORDS_TH = ["รายงานการศึกษาอิสระ", "มหาวิทยาลัยขอนแก่น"]
SLOT1_REQUIRED_KEYWORDS_EN = ["INDEPENDENT STUDY", "KHON KAEN UNIVERSITY"]


def check_slot1_keywords(doc: Document) -> list[HardRuleError]:
    full_text = "\n".join(p.text for p in doc.paragraphs).upper()
    errors = []
    for kw in SLOT1_REQUIRED_KEYWORDS_TH:
        if kw not in "\n".join(p.text for p in doc.paragraphs):
            errors.append(HardRuleError(
                location="Cover Page",
                issue=f"ไม่พบคำสำคัญ: '{kw}'",
                detail=f"ปกนอก/ปกในควรมีข้อความ '{kw}'",
                severity="warning",
                rule_type="keyword",
            ))
    for kw in SLOT1_REQUIRED_KEYWORDS_EN:
        if kw not in full_text:
            errors.append(HardRuleError(
                location="Cover Page",
                issue=f"Missing keyword: '{kw}'",
                detail=f"Cover page should contain '{kw}'",
                severity="warning",
                rule_type="keyword",
            ))
    return errors


def check_slot1_title_size(doc: Document) -> list[HardRuleError]:
    errors = []
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        found_sizes = []
        is_bold = False
        for run in para.runs:
            if run.font.size:
                found_sizes.append(run.font.size.pt)
            if run.font.bold:
                is_bold = True
        has_valid = any(s in (16, 18) for s in found_sizes)
        sz_str = ", ".join(f"{s:.0f}pt" for s in found_sizes) if found_sizes else "ไม่ระบุ"
        preview = para.text.strip()[:50]
        if found_sizes and not has_valid:
            errors.append(HardRuleError(
                location=f"Paragraph {i+1} (Title)",
                issue="ขนาดหัวข้อผิด (Wrong Title Size)",
                detail=f"หัวข้อ: \"{preview}\" — พบ {sz_str} (ควรเป็น 16pt หรือ 18pt ตัวหนา)",
                severity="error",
                rule_type="title_size",
            ))
        elif found_sizes and has_valid and not is_bold:
            errors.append(HardRuleError(
                location=f"Paragraph {i+1} (Title)",
                issue="หัวข้อไม่เป็นตัวหนา (Title Not Bold)",
                detail=f"หัวข้อ: \"{preview}\" — ขนาดถูกต้อง ({sz_str}) แต่ต้องเป็นตัวหนา",
                severity="warning",
                rule_type="title_size",
            ))
        break
    return errors


# ──────────────────────────────────────────────
# Logo validation (Slot 1 only)
# ──────────────────────────────────────────────

def check_logo(doc: Document) -> list[HardRuleError]:
    errors = []
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            images.append(rel.target_part.blob)
    if not images:
        errors.append(HardRuleError(
            location="Cover Page",
            issue="ไม่พบตรามหาวิทยาลัย (Logo Missing)",
            detail="ไม่พบรูปภาพตรามหาวิทยาลัยในเอกสาร ควรมีตราอยู่ที่หน้าปก",
            severity="error",
            rule_type="logo",
        ))
        return errors
    if not HAS_PIL:
        return errors
    try:
        img = Image.open(io.BytesIO(images[0]))
        if img.mode not in ("L", "LA", "1"):
            errors.append(HardRuleError(
                location="Cover Page, Image 1",
                issue="ตราเป็นภาพสี (Logo Not Grayscale)",
                detail=f"ตรามหาวิทยาลัยเป็นภาพโหมด '{img.mode}' ควรเป็นภาพขาวดำ",
                severity="error",
                rule_type="logo",
            ))
    except Exception:
        pass
    return errors


# ──────────────────────────────────────────────
# Slot 2: Abstract checks
# ──────────────────────────────────────────────

SLOT2_TITLES = ["บทคัดย่อ", "ABSTRACT"]


def check_slot2_indentation(doc: Document) -> list[HardRuleError]:
    """Check 0.5-inch first-line indent on body paragraphs."""
    errors = []
    checked = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or len(text) < 20:
            continue
        if any(t in text for t in SLOT2_TITLES):
            continue
        pf = para.paragraph_format
        indent = pf.first_line_indent
        if indent is not None:
            indent_in = indent.inches
            if abs(indent_in - 0.5) > 0.15:
                errors.append(HardRuleError(
                    location=f"Paragraph {i+1}",
                    issue="การเยื้องย่อหน้าผิด (Wrong Indentation)",
                    detail=(
                        f"ย่อหน้าเยื้อง {indent_in:.2f} นิ้ว "
                        f"(ควรเป็น 0.5 นิ้ว / ประมาณ 7 ตัวอักษร)"
                    ),
                    severity="warning",
                    rule_type="indentation",
                ))
        checked += 1
        if checked >= 5:
            break
    return errors


# ──────────────────────────────────────────────
# Slot 3: TOC checks
# ──────────────────────────────────────────────

SLOT3_TITLES = ["กิตติกรรมประกาศ", "สารบัญ", "สารบัญตาราง", "สารบัญภาพ"]
THAI_PAGE_NUMS = set("กขคฅฆงจฉชซฌญฎฏฐฑฒณดตถทธนบปผฝพฟภมยรลวศษสหฬอฮ")


def check_slot3_thai_page_numbers(doc: Document) -> list[HardRuleError]:
    """Check if TOC pages use Thai alphabetic page numbering."""
    errors = []
    full_text = "\n".join(p.text for p in doc.paragraphs)
    has_thai_nums = any(c in THAI_PAGE_NUMS for c in full_text)
    if not has_thai_nums:
        errors.append(HardRuleError(
            location="Table of Contents",
            issue="ไม่พบเลขหน้าแบบอักษรไทย (Missing Thai Page Numbers)",
            detail="สารบัญควรใช้เลขหน้าแบบ ก, ข, ค... ไม่ใช่เลขอารบิก",
            severity="warning",
            rule_type="page_number",
        ))
    return errors


# ──────────────────────────────────────────────
# Slot 4: Main Content checks
# ──────────────────────────────────────────────

CHAPTER_TITLE_PATTERN = re.compile(r'^บทที่\s*\d+', re.UNICODE)


def check_slot4_chapter_title(doc: Document) -> list[HardRuleError]:
    errors = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if CHAPTER_TITLE_PATTERN.match(text):
            sizes = [r.font.size.pt for r in para.runs if r.font.size]
            has_18 = any(abs(s - 18) < 0.5 for s in sizes)
            if sizes and not has_18:
                sz_str = ", ".join(f"{s:.0f}pt" for s in sizes)
                errors.append(HardRuleError(
                    location=f"Paragraph {i+1}",
                    issue=f"ขนาดชื่อบทผิด (Chapter Title ≠ 18pt)",
                    detail=f"\"{text[:40]}\" — พบ {sz_str} (ควรเป็น 18pt)",
                    severity="error",
                    rule_type="title_size",
                ))
            alignment = para.alignment
            if alignment is not None and alignment != WD_ALIGN_PARAGRAPH.CENTER:
                errors.append(HardRuleError(
                    location=f"Paragraph {i+1}",
                    issue="ชื่อบทไม่จัดกึ่งกลาง (Chapter Title Not Centered)",
                    detail=f"\"{text[:40]}\" ควรจัดกึ่งกลาง",
                    severity="warning",
                    rule_type="title_align",
                ))
            break
    return errors


# ──────────────────────────────────────────────
# Slot 5: Bibliography / Appendix checks
# ──────────────────────────────────────────────

SLOT5_TITLES = ["เอกสารอ้างอิง", "บรรณานุกรม", "ภาคผนวก", "REFERENCES", "BIBLIOGRAPHY", "APPENDIX"]


def check_slot5_hanging_indent(doc: Document) -> list[HardRuleError]:
    """APA 7th: second line of citations should have hanging indent."""
    errors = []
    checked = 0
    in_refs = False
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if any(t in text.upper() for t in ["เอกสารอ้างอิง", "บรรณานุกรม", "REFERENCES", "BIBLIOGRAPHY"]):
            in_refs = True
            continue
        if any(t in text.upper() for t in ["ภาคผนวก", "APPENDIX"]):
            in_refs = False
            continue
        if not in_refs:
            continue
        if len(text) < 15:
            continue
        pf = para.paragraph_format
        hanging = None
        if pf.first_line_indent and pf.first_line_indent.inches < 0:
            hanging = abs(pf.first_line_indent.inches)
        if hanging is None or hanging < 0.2:
            errors.append(HardRuleError(
                location=f"Paragraph {i+1}",
                issue="ไม่พบ Hanging Indent ในรายการอ้างอิง",
                detail=(
                    f"รายการอ้างอิง: \"{text[:50]}...\" — "
                    f"APA 7th กำหนดให้บรรทัดที่ 2 เป็นต้นไปต้องเยื้อง (Hanging Indent ~0.5 นิ้ว)"
                ),
                severity="warning",
                rule_type="hanging_indent",
            ))
        checked += 1
        if checked >= 8:
            break
    return errors


# ──────────────────────────────────────────────
# Extract full text from document
# ──────────────────────────────────────────────

def extract_text(doc: Document) -> str:
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


# ──────────────────────────────────────────────
# TOC Heading Extraction (Slot 3)
# ──────────────────────────────────────────────

_TOC_STRIP_RE = re.compile(r'[\.\ \-–—\s\t]+\d+\s*$')

def extract_toc_headings(doc: Document) -> list[str]:
    headings = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) < 3:
            continue
        clean = _TOC_STRIP_RE.sub('', text).strip()
        clean = clean.replace('\t', ' ').strip()
        if clean and len(clean) >= 3:
            headings.append(clean)
    return headings


# ──────────────────────────────────────────────
# Content Heading Extraction (Slot 4 chapters)
# ──────────────────────────────────────────────

def extract_content_headings(doc: Document) -> list[str]:
    headings = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) < 3:
            continue
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            headings.append(text)
            continue
        runs_with_text = [r for r in para.runs if r.text.strip()]
        if runs_with_text and all(r.bold for r in runs_with_text):
            if len(text) < 120:
                headings.append(text)
    return headings


# ──────────────────────────────────────────────
# TOC vs Content Matcher
# ──────────────────────────────────────────────

def match_toc_vs_content(
    toc_headings: list[str],
    content_headings: list[str],
) -> list[HardRuleError]:
    if not toc_headings:
        return []
    errors = []
    toc_set = {h.strip() for h in toc_headings}
    for heading in content_headings:
        clean = heading.strip()
        if clean not in toc_set:
            errors.append(HardRuleError(
                location="Content Heading",
                issue="หัวข้อไม่ตรงกับสารบัญ (TOC Mismatch)",
                detail=(
                    f"หัวข้อ \"{clean[:80]}\" ในเนื้อหาไม่ตรงกับรายการในสารบัญ "
                    f"— กรุณาตรวจสอบการสะกดและอัปเดตสารบัญให้ตรงกัน"
                ),
                severity="warning",
                rule_type="toc_mismatch",
            ))
    return errors


# ──────────────────────────────────────────────
# SLOT_NAMES
# ──────────────────────────────────────────────

SLOT_NAMES = {
    1: "ส่วนหน้า (Front Matter)",
    2: "บทคัดย่อ (Abstract)",
    3: "สารบัญ (Table of Contents)",
    4: "เนื้อหา (Main Content)",
    5: "บรรณานุกรม (References)",
}


# ──────────────────────────────────────────────
# Main entry point per slot — ALL SLOTS COVERED
# ──────────────────────────────────────────────

def process_slot(doc: Document, slot_number: int) -> list[HardRuleError]:
    """
    Run all applicable hard-rule checks for the given slot.
    Returns a combined list of HardRuleError with location/issue/detail.
    """
    errors = []

    # ── Global checks (all slots) ──
    errors.extend(check_margins(doc, slot_number))
    errors.extend(check_fonts(doc))

    # ── Slot-specific checks ──
    if slot_number == 1:
        errors.extend(check_slot1_title_size(doc))
        errors.extend(check_logo(doc))
        errors.extend(check_slot1_keywords(doc))

    elif slot_number == 2:
        errors.extend(check_titles(doc, SLOT2_TITLES, "Abstract"))
        errors.extend(check_slot2_indentation(doc))
        errors.extend(check_normal_text_size(doc))

    elif slot_number == 3:
        errors.extend(check_titles(doc, SLOT3_TITLES, "TOC"))
        errors.extend(check_slot3_thai_page_numbers(doc))

    elif slot_number == 4:
        errors.extend(check_slot4_chapter_title(doc))
        errors.extend(check_normal_text_size(doc))

    elif slot_number == 5:
        errors.extend(check_titles(doc, SLOT5_TITLES, "References"))
        errors.extend(check_slot5_hanging_indent(doc))
        errors.extend(check_normal_text_size(doc))

    return errors
